from flask import Flask, render_template, request, redirect, url_for, send_file
import os
import pandas as pd
from docx import Document


app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_data_from_word(doc_path):
    # Read the Word document
    doc = Document(doc_path)

    # Extract lines 2, 4, 5, 6, and the paragraph starting with "ABSTRACT"
    lines = [
        doc.paragraphs[1].text,
        doc.paragraphs[3].text,
        doc.paragraphs[4].text,
        doc.paragraphs[5].text,
        ""
    ]

    abstract_started = False
    for paragraph in doc.paragraphs:
        if abstract_started:
            # Stop when encountering the last line in the "ABSTRACT" paragraph
            if paragraph.text.strip() == "":
                break
            lines.append(paragraph.text)
        elif paragraph.text.strip().upper() == "ABSTRACT":
            # Start extracting when reaching the "ABSTRACT" heading
            abstract_started = True

    return lines

def process_folder(file_list, output_excel):
    # Check if the output Excel file exists
    if os.path.exists(output_excel):
        # Load existing data from the Excel file
        df_existing = pd.read_excel(output_excel)
    else:
        # Create an empty DataFrame if the Excel file doesn't exist
        df_existing = pd.DataFrame(columns=["TITLE", "STUDENT NAME", "MATRIC NUMBER", "SUPERVISOR", "ABSTRACT"])

    # Process each uploaded Word document
    for file in file_list:
        doc_path = os.path.join("uploads", file.filename)
        file.save(doc_path)

        # Extract data from the Word document
        lines = extract_data_from_word(doc_path)

        # Create a new DataFrame with the extracted lines
        df_new = pd.DataFrame([lines], columns=["TITLE", "STUDENT NAME", "MATRIC NUMBER", "SUPERVISOR", "", "ABSTRACT"])

        # Concatenate the new DataFrame with the existing one
        df_existing = pd.concat([df_existing, df_new], ignore_index=True)

    # Save the combined DataFrame to the output Excel file
    df_existing.to_excel(output_excel, index=False, engine='openpyxl')


from flask import Flask, render_template, request, redirect, url_for, send_file
import os
import pandas as pd
from docx import Document

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ... (Rest of the code remains unchanged)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        files = request.files.getlist('file')
        if files and all(allowed_file(file.filename) for file in files):
            process_folder(files, "Extract.xlsx")
            return redirect(url_for('success'))

    return render_template('index.html')



@app.route('/success')
def success():
    return render_template('success.html')

@app.route('/download')
def download():
    path = "Extract.xlsx"
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=os.environ.get('PORT', 5000))
